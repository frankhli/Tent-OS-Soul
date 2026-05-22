"""Word 渲染引擎 —— python-docx

将 WordDocument 数据结构渲染为 .docx 文件。
支持：段落、表格、图片、页眉页脚、目录、样式。
"""

import base64
import io
import re
from pathlib import Path
from typing import Dict, Any
from urllib.parse import urlparse

import httpx

from tent_os.skills.word.schema import WordDocument, WordParagraph, WordTable, WordImage, WordHeaderFooter


class WordRenderer:
    """Word 文档渲染器"""
    
    def render(self, doc: WordDocument, output_path: str) -> str:
        """渲染 WordDocument 为 .docx 文件"""
        try:
            from docx import Document as DocxDocument
            from docx.shared import Pt, Mm, Cm, RGBColor, Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.enum.style import WD_STYLE_TYPE
        except ImportError:
            raise ImportError("python-docx 未安装，请运行: pip install python-docx")
        
        output = Path(output_path)
        output.parent.mkdir(parents=True, exist_ok=True)
        
        # 创建文档
        document = DocxDocument()
        
        cfg = doc.config
        font_name = cfg.get("font_name", "微软雅黑")
        font_size_pt = cfg.get("font_size", 10.5)
        line_spacing = cfg.get("line_spacing", 1.5)
        
        # 设置默认字体
        style = document.styles["Normal"]
        font = style.font
        font.name = font_name
        font.size = Pt(font_size_pt)
        
        # 设置中文字体（通过 eastAsia）
        rFonts = style.element.rPr.rFonts if style.element.rPr is not None else None
        if rFonts is not None:
            rFonts.set("{http://schemas.openxmlformats.org/drawingml/2006/main}eastAsia", font_name)
        
        # 设置页面边距
        sections = document.sections[0]
        sections.top_margin = Mm(cfg.get("page_margin_top", 25))
        sections.bottom_margin = Mm(cfg.get("page_margin_bottom", 25))
        sections.left_margin = Mm(cfg.get("page_margin_left", 30))
        sections.right_margin = Mm(cfg.get("page_margin_right", 30))
        
        # 页眉页脚
        if doc.header_footer:
            self._apply_header_footer(document, doc.header_footer)
        
        # 标题页
        if doc.title:
            p = document.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(doc.title)
            run.bold = True
            run.font.size = Pt(22)
            run.font.name = font_name
            p.space_after = Pt(12)
        
        if doc.subtitle:
            p = document.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(doc.subtitle)
            run.italic = True
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
            run.font.name = font_name
            p.space_after = Pt(24)
        
        if doc.author or doc.date:
            meta_text = ""
            if doc.author:
                meta_text += f"作者: {doc.author}"
            if doc.date:
                meta_text += f"   日期: {doc.date}" if meta_text else f"日期: {doc.date}"
            p = document.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(meta_text)
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
            run.font.name = font_name
            p.space_after = Pt(36)
        
        # 目录
        if cfg.get("show_toc", False):
            toc_para = document.add_paragraph()
            toc_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = toc_para.add_run("目录")
            run.bold = True
            run.font.size = Pt(16)
            run.font.name = font_name
            toc_para.space_after = Pt(12)
        
        # 渲染内容块
        for block in doc.blocks:
            b_type = block.get("type", "paragraph")
            b_data = block.get("data", {})
            
            if b_type == "paragraph":
                para = b_data if isinstance(b_data, WordParagraph) else WordParagraph(**b_data)
                self._render_paragraph(document, para, font_name, line_spacing)
            elif b_type == "table":
                table = b_data if isinstance(b_data, WordTable) else WordTable(**b_data)
                self._render_table(document, table, font_name)
            elif b_type == "image":
                img = b_data if isinstance(b_data, WordImage) else WordImage(**b_data)
                self._render_image(document, img)
        
        document.save(str(output))
        return str(output)
    
    def _render_paragraph(self, document, para: WordParagraph, font_name: str, line_spacing: float):
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        # 样式映射
        style_map = {
            "Heading 1": "Heading 1",
            "Heading 2": "Heading 2",
            "Heading 3": "Heading 3",
            "Quote": "Quote",
            "List Bullet": "List Bullet",
            "List Number": "List Number",
            "Normal": "Normal",
        }
        docx_style = style_map.get(para.style, "Normal")
        
        p = document.add_paragraph(style=docx_style)
        
        # 分页符
        if para.page_break_before:
            p.paragraph_format.page_break_before = True
        
        # 行距
        p.paragraph_format.line_spacing = line_spacing
        
        # 对齐
        align_map = {
            "left": WD_ALIGN_PARAGRAPH.LEFT,
            "center": WD_ALIGN_PARAGRAPH.CENTER,
            "right": WD_ALIGN_PARAGRAPH.RIGHT,
            "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
        }
        if para.alignment:
            p.alignment = align_map.get(para.alignment, WD_ALIGN_PARAGRAPH.LEFT)
        
        # 文本内容
        run = p.add_run(para.text)
        run.font.name = font_name
        if para.bold:
            run.bold = True
        if para.italic:
            run.italic = True
        if para.font_size:
            run.font.size = Pt(para.font_size)
        if para.font_color:
            color = para.font_color.lstrip("#")
            if len(color) == 6:
                run.font.color.rgb = RGBColor(int(color[0:2], 16), int(color[2:4], 16), int(color[4:6], 16))
    
    def _render_table(self, document, table: WordTable, font_name: str):
        from docx.shared import Mm, Pt, RGBColor
        
        rows = 1 + len(table.rows) if table.headers else len(table.rows)
        cols = max(len(table.headers) if table.headers else 0, max((len(r) for r in table.rows), default=0))
        if rows == 0 or cols == 0:
            return
        
        docx_table = document.add_table(rows=rows, cols=cols)
        docx_table.style = table.style if table.style else "Table Grid"
        
        # 设置列宽
        for idx, width in enumerate(table.column_widths):
            if idx < cols:
                for cell in docx_table.columns[idx].cells:
                    cell.width = Mm(width)
        
        row_idx = 0
        
        # 表头
        if table.headers:
            for col_idx, header in enumerate(table.headers):
                if col_idx < cols:
                    cell = docx_table.rows[row_idx].cells[col_idx]
                    cell.text = str(header)
                    # 表头加粗
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
                            run.font.name = font_name
            row_idx += 1
        
        # 数据行
        for data_row in table.rows:
            for col_idx, value in enumerate(data_row):
                if col_idx < cols:
                    cell = docx_table.rows[row_idx].cells[col_idx]
                    cell.text = str(value)
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.name = font_name
            row_idx += 1
    
    def _render_image(self, document, img: WordImage):
        from docx.shared import Mm
        
        image_bytes = None
        source = img.source
        
        # 尝试获取图片数据
        if source.startswith("http"):
            try:
                resp = httpx.get(source, timeout=10)
                if resp.status_code == 200:
                    image_bytes = io.BytesIO(resp.content)
            except Exception:
                pass
        elif source.startswith("data:image"):
            # base64 data URI
            try:
                match = re.match(r"data:image/\w+;base64,(.+)", source)
                if match:
                    image_bytes = io.BytesIO(base64.b64decode(match.group(1)))
            except Exception:
                pass
        
        if image_bytes:
            kwargs = {}
            if img.width:
                kwargs["width"] = Mm(img.width)
            if img.height:
                kwargs["height"] = Mm(img.height)
            document.add_picture(image_bytes, **kwargs)
            
            if img.caption:
                cap = document.add_paragraph()
                cap.alignment = 1  # center
                run = cap.add_run(img.caption)
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    
    def _apply_header_footer(self, document, hf: WordHeaderFooter):
        from docx.shared import Pt, RGBColor
        
        section = document.sections[0]
        
        if hf.header_text:
            header = section.header
            header.is_linked_to_previous = False
            p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
            p.text = hf.header_text
            for run in p.runs:
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
        
        if hf.footer_text or hf.show_page_numbers:
            footer = section.footer
            footer.is_linked_to_previous = False
            p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
            p.alignment = 1  # center
            
            if hf.footer_text:
                run = p.add_run(hf.footer_text)
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)


def render_word(doc: WordDocument, output_path: str) -> str:
    """便捷函数：渲染 Word 并保存"""
    renderer = WordRenderer()
    return renderer.render(doc, output_path)
